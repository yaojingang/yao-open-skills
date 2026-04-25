#!/usr/bin/env python3
from __future__ import annotations

import argparse
import importlib.util
import json
import math
import os
import sys
from html import escape
from pathlib import Path
from typing import Any


def maybe_reexec_with_bundled_python() -> None:
    missing = [name for name in ("docx", "reportlab") if importlib.util.find_spec(name) is None]
    if not missing:
        return
    if os.environ.get("CODEX_BUNDLED_PYTHON_REEXEC") == "1":
        raise SystemExit(f"Missing Python packages after re-exec: {', '.join(missing)}")

    candidates = [
        os.environ.get("CODEX_PRIMARY_RUNTIME_PYTHON"),
        str(Path.home() / ".cache" / "codex-runtimes" / "codex-primary-runtime" / "dependencies" / "python" / "bin" / "python3"),
    ]
    for candidate in candidates:
        if not candidate:
            continue
        path = Path(candidate)
        if not path.exists():
            continue
        if path.resolve() == Path(sys.executable).resolve():
            continue
        env = os.environ.copy()
        env["CODEX_BUNDLED_PYTHON_REEXEC"] = "1"
        os.execve(str(path), [str(path), __file__, *sys.argv[1:]], env)

    raise SystemExit(f"Unable to locate a bundled Python runtime with: {', '.join(missing)}")


maybe_reexec_with_bundled_python()

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from bayesian_decision_report import build_report, load_request, localize_text


ROOT = Path(__file__).resolve().parent.parent
CSS_PATH = ROOT / "templates" / "report-theme.css"

TOP_NAV = [
    ("summary", "先看结论", "Summary"),
    ("action-plan", "行动建议", "Action"),
    ("why", "为什么这样判断", "Why"),
    ("conversation", "对话过程", "Conversation"),
    ("decision", "问题定义", "Decision"),
    ("evidence", "证据", "Evidence"),
    ("update", "概率判断", "Update"),
    ("prior", "先验", "Prior"),
    ("sensitivity", "敏感性", "Sensitivity"),
    ("warnings", "注意事项", "Warnings"),
    ("workflow", "附录", "Appendix"),
]

SKILL_FLOW = [
    {"zh": "问题结构化", "en": "Structure the decision"},
    {"zh": "基于当前信息给初步先验", "en": "Set an initial prior from the current information"},
    {"zh": "通过多轮追问补齐关键缺口", "en": "Close the key gaps through iterative questioning"},
    {"zh": "检索与分级证据", "en": "Collect and grade evidence"},
    {"zh": "构造先验与参考类", "en": "Build priors and reference classes"},
    {"zh": "执行贝叶斯更新", "en": "Run the Bayesian update"},
    {"zh": "做敏感性与阈值分析", "en": "Stress-test thresholds and sensitivity"},
    {"zh": "给出行动与下一步信息建议", "en": "Recommend the next action and information step"},
    {"zh": "记录每轮变化并输出过程日志", "en": "Record each round and export a process log"},
]

SKILL_CAPABILITIES = [
    {
        "zh_title": "问题定义与参考类选择",
        "zh_body": "把模糊问题压缩成可计算假设、时间范围、成功标准和行动集合，并显式选择参考类。",
        "en_title": "Problem framing and reference classes",
        "en_body": "Turns a fuzzy choice into a computable hypothesis, time horizon, success metric, and action set with an explicit reference class.",
    },
    {
        "zh_title": "证据分级与先验构造",
        "zh_body": "区分强证据、弱证据和类比信号，输出带来源、可信度和等效样本量的先验。",
        "en_title": "Evidence grading and prior construction",
        "en_body": "Separates strong evidence from weak signals and produces priors with source notes, quality, and equivalent sample size.",
    },
    {
        "zh_title": "赔率更新与共轭更新",
        "zh_body": "支持似然比赔率更新、Beta-binomial 计数更新，以及二者的组合路径。",
        "en_title": "Odds updates and conjugate updates",
        "en_body": "Supports likelihood-ratio odds updates, Beta-binomial count updates, and the combined path when both are needed.",
    },
    {
        "zh_title": "行动阈值与期望值比较",
        "zh_body": "把后验概率映射到行动阈值和期望值排序，避免只停留在“概率高低”的描述。",
        "en_title": "Action thresholds and expected value",
        "en_body": "Converts posterior beliefs into action thresholds and expected-value ranking instead of stopping at raw probability.",
    },
    {
        "zh_title": "敏感性分析与防幻觉提醒",
        "zh_body": "用不同先验、证据强度和依赖折扣重跑结果，显式标记结论是否稳健。",
        "en_title": "Sensitivity analysis and fragility checks",
        "en_body": "Re-runs the decision under alternative priors, evidence strength, and dependency discounts to show whether the conclusion is robust.",
    },
    {
        "zh_title": "自动化报告 bundle",
        "zh_body": "从同一个结构化输入自动生成 JSON、Markdown、双语 HTML、PDF 和 Word 报告，并保持内容同步。",
        "en_title": "Automated report bundle",
        "en_body": "Automatically generates synchronized JSON, Markdown, bilingual HTML, PDF, and Word reports from the same structured input.",
    },
    {
        "zh_title": "多轮对话决策循环",
        "zh_body": "先基于当前信息给弱先验和初步判断，再通过多轮对话追问缺失信息，持续更新先验、后验和决策准备度。",
        "en_title": "Multi-turn decision loop",
        "en_body": "Starts with a weak prior and an initial read, then iteratively asks for missing information and updates the prior, posterior, and decision readiness over multiple turns.",
    },
    {
        "zh_title": "过程日志与变化图表",
        "zh_body": "自动记录每一轮对话里新增了什么信息、如何完成贝叶斯更新、判断为什么改变，并在报告中输出轨迹图和过程表。",
        "en_title": "Process log and change chart",
        "en_body": "Records what changed in each round, how the Bayesian update was applied, why the judgment moved, and renders both a trajectory chart and a process table in the report.",
    },
]

WARNING_TRANSLATIONS = {
    "未提供先验等效样本量，先验强度相关的敏感性判断会更弱。": "No prior equivalent sample size was provided, so sensitivity about prior strength is weaker.",
    "在 Beta-binomial 更新之后又追加了证据更新，请确认这些证据没有被重复包含在已观测计数中。": "Additional evidence was applied after the Beta-binomial update. Confirm that these signals are not already embedded in the observed counts.",
    "部分证据较弱或较间接，后验判断应视为更脆弱。": "Some evidence is weak or indirect, so the posterior judgment should be treated as more fragile.",
    "至少有一条证据因依赖关系被折扣处理，不应把这次更新理解为完全独立证据的简单叠加。": "At least one evidence item was discounted for dependence, so this update should not be read as a simple stack of fully independent signals.",
    "本报告仅用于辅助决策，不替代持证专业人士的正式判断。": "This report is decision support only and does not replace judgment from a licensed professional.",
    "这里的后验区间来自赔率更新下的敏感性包络，并非严格统计意义上的可信区间。": "The posterior interval here is a sensitivity envelope from odds updates, not a strict statistical credible interval.",
}

CONFIDENCE_EN = {
    "low": "Low",
    "medium": "Medium",
    "medium-high": "Medium-high",
}

STABILITY_EN = {
    "stable": "Stable",
    "mixed": "Mixed",
    "unstable": "Unstable",
}

READINESS_EN = {
    "可以决策": "Ready to decide",
    "接近可决策": "Nearly ready",
    "仍需补信息": "Need more information",
    "未评估": "Not evaluated",
}

DECISION_STATUS_EN = {
    "ready": "Ready to decide",
    "needs-more-info": "Need more information",
    "in-progress": "Keep iterating",
    "blocked": "Blocked",
}

DIRECTION_EN = {
    "support": "Supports",
    "against": "Weakens",
}

DECISION_STATES = {
    "proceed": {
        "label_zh": "可以推进",
        "label_en": "Proceed",
        "tone_zh": "当前判断偏向正面，可以往前走，但仍建议按推荐节奏推进。",
        "tone_en": "The current read is positive enough to move forward, but you should still follow the recommended pace.",
    },
    "confirm": {
        "label_zh": "先确认，再推进",
        "label_en": "Confirm first",
        "tone_zh": "现在不是不能做，而是还差最后一两个关键信息，先确认再推进更稳。",
        "tone_en": "This is not a no. It means one or two key inputs are still missing, so confirm them before moving forward.",
    },
    "hold": {
        "label_zh": "暂时别做",
        "label_en": "Hold off",
        "tone_zh": "当前把握还不够，先暂缓更稳，等拿到更强证据再决定。",
        "tone_en": "The current case is not strong enough, so holding off is safer until stronger evidence appears.",
    },
}


def request_text(value: Any, lang: str) -> str:
    text = localize_text(value, lang)
    return text.strip() if isinstance(text, str) and text.strip() else "-"


def localize_items(values: list[Any] | None, lang: str) -> list[str]:
    items = []
    for value in values or []:
        text = request_text(value, lang)
        if text != "-":
            items.append(text)
    return items


def fmt_pct(value: Any) -> str:
    if value is None:
        return "-"
    return f"{float(value) * 100:.1f}%"


def fmt_num(value: Any) -> str:
    if value is None:
        return "-"
    number = float(value)
    if abs(number - round(number)) < 1e-9:
        return f"{int(round(number)):,}"
    return f"{number:,.1f}"


def fmt_interval(interval: list[float] | None) -> str:
    if not interval:
        return "-"
    return f"{fmt_pct(interval[0])} - {fmt_pct(interval[1])}"


def html_text(value: Any) -> str:
    return escape(str(value))


def dual_html(zh: str, en: str) -> str:
    zh_text = html_text(zh or "-")
    en_text = html_text(en or zh or "-")
    return f'<span class="lang-zh">{zh_text}</span><span class="lang-en">{en_text}</span>'


def dual_tag(tag: str, zh: str, en: str, class_name: str = "") -> str:
    attrs = f' class="{class_name}"' if class_name else ""
    return f"<{tag}{attrs}>{dual_html(zh, en)}</{tag}>"


def fold_section(section_id: str, zh_title: str, en_title: str, body_html: str, kicker_zh: str = "默认折叠", kicker_en: str = "Collapsed by default") -> str:
    return f"""
    <details class="fold-section pro-only section-anchor-offset" id="{section_id}">
      <summary>
        <div class="fold-heading">
          <div class="section-kicker">{dual_html(kicker_zh, kicker_en)}</div>
          <div class="fold-title">{dual_html(zh_title, en_title)}</div>
        </div>
        <span class="fold-indicator">{dual_html("展开", "Open")}</span>
      </summary>
      <div class="fold-panel">
        {body_html}
      </div>
    </details>
    """


def action_name_maps(request: dict) -> tuple[dict[str, str], dict[str, str]]:
    zh_to_en: dict[str, str] = {}
    en_to_zh: dict[str, str] = {}
    for action in request.get("actions", []):
        zh_name = request_text(action.get("name"), "zh")
        en_name = request_text(action.get("name"), "en")
        zh_to_en[zh_name] = en_name
        en_to_zh[en_name] = zh_name
    return zh_to_en, en_to_zh


def recommended_action_en(request: dict, recommended_zh: str) -> str:
    zh_to_en, _ = action_name_maps(request)
    return zh_to_en.get(recommended_zh, recommended_zh)


def summary_en(request: dict, report: dict) -> str:
    summary = report["summary"]
    action_en = recommended_action_en(request, summary["recommendation"])
    return f"The current posterior probability is about {fmt_pct(summary['posterior_probability'])}; given the available evidence, the better action is {action_en}."


def prior_confidence_en(request: dict, report: dict) -> str:
    raw = request.get("prior", {}).get("confidence")
    text = request_text(raw, "en")
    if text != "-":
        return text.replace("-", " ").title()
    code = report["summary"].get("confidence_code")
    return CONFIDENCE_EN.get(code, code or "-")


def stability_en(report: dict) -> str:
    code = report["sensitivity"].get("conclusion_stability_code")
    return STABILITY_EN.get(code, code or "-")


def direction_en(item: dict) -> str:
    code = item.get("direction_code")
    return DIRECTION_EN.get(code, code or "-")


def warning_en(text: str) -> str:
    return WARNING_TRANSLATIONS.get(text, text)


def natural_frequency_en(report: dict) -> str:
    natural = report["natural_frequency"]
    base = natural["base_population"]
    prior_cases = natural["prior_expected_successes"]
    posterior_cases = natural["posterior_expected_successes"]
    return (
        f"Out of {base} similar cases, the prior suggests about {prior_cases} would succeed. "
        f"After the current evidence, about {posterior_cases} would succeed."
    )


def next_information_reason_en(report: dict) -> str:
    if report["next_information"].get("high_value"):
        return "The current result is close to an action threshold, so one more strong signal could flip the recommendation."
    return "More information still helps, but the current recommendation is less sensitive to boundary changes."


def first_next_information(request: dict, lang: str) -> str:
    candidates = (request.get("next_information") or {}).get("candidates") or []
    if not candidates:
        return "-"
    return request_text(candidates[0], lang)


def decision_reason_en(report: dict) -> str:
    if report["decision"]["expected_value_ranking"]:
        return "Expected-value comparison based on the current posterior places this action first."
    return "No action-utility inputs were provided, so the skill cannot rank actions directly."


def conversation_raw_rounds(request: dict) -> list[dict]:
    return ((request.get("conversation") or {}).get("rounds") or [])


def conversation_round_en(raw_round: dict, key: str, default: str = "-") -> str:
    value = request_text(raw_round.get(key), "en")
    return value if value != "-" else default


def readiness_en(label_zh: str) -> str:
    return READINESS_EN.get(label_zh, label_zh)


def decision_status_en(report: dict) -> str:
    process = report.get("conversation_process") or {}
    code = process.get("status_code")
    return DECISION_STATUS_EN.get(code, process.get("status_en") or code or "-")


def conversation_dual_list_items(round_payload: dict, raw_round: dict, primary_key: str, fallback_key: str | None = None) -> str:
    zh_items = list(round_payload.get(primary_key) or [])
    if not zh_items and fallback_key:
        zh_items = list(round_payload.get(fallback_key) or [])
    raw_items = list(raw_round.get(primary_key) or [])
    if not raw_items and fallback_key:
        raw_items = list(raw_round.get(fallback_key) or [])
    count = max(len(zh_items), len(raw_items))
    if count == 0:
        return ""
    items = []
    for index in range(count):
        zh_value = zh_items[index] if index < len(zh_items) else "-"
        en_source = raw_items[index] if index < len(raw_items) else zh_value
        items.append(f"<li>{dual_html(zh_value, request_text(en_source, 'en'))}</li>")
    return "".join(items)


def conversation_chart_svg(report: dict) -> str:
    process = report.get("conversation_process") or {}
    points = process.get("trajectory") or []
    if not points:
        return ""

    width = 760
    height = 260
    margin_left = 48
    margin_right = 20
    margin_top = 20
    margin_bottom = 42
    inner_width = width - margin_left - margin_right
    inner_height = height - margin_top - margin_bottom
    count = len(points)
    step = inner_width / max(count - 1, 1)

    def x_at(index: int) -> float:
        return margin_left + step * index

    def y_at(value: float | None) -> float:
        safe = 0.0 if value is None else max(0.0, min(1.0, float(value)))
        return margin_top + (1.0 - safe) * inner_height

    def polyline(values: list[float | None]) -> str:
        coords = [f"{x_at(index):.1f},{y_at(value):.1f}" for index, value in enumerate(values) if value is not None]
        return " ".join(coords)

    grid_lines = []
    for ratio in (0.0, 0.25, 0.5, 0.75, 1.0):
        y = y_at(ratio)
        grid_lines.append(f'<line x1="{margin_left}" y1="{y:.1f}" x2="{width - margin_right}" y2="{y:.1f}" class="chart-grid" />')
        grid_lines.append(
            f'<text x="{margin_left - 10}" y="{y + 4:.1f}" text-anchor="end" class="chart-axis-label">{ratio * 100:.0f}%</text>'
        )

    prior_values = [item.get("prior_probability") for item in points]
    posterior_values = [item.get("posterior_probability") for item in points]
    readiness_values = [item.get("decision_readiness") for item in points]

    readiness_bars = []
    point_markers = []
    for index, point in enumerate(points):
        x = x_at(index)
        readiness = point.get("decision_readiness")
        if readiness is not None:
            y = y_at(readiness)
            height_value = margin_top + inner_height - y
            readiness_bars.append(
                f'<rect x="{x - 18:.1f}" y="{y:.1f}" width="36" height="{height_value:.1f}" class="chart-readiness-bar" rx="8" />'
            )
        for value, class_name in (
            (point.get("prior_probability"), "chart-point-prior"),
            (point.get("posterior_probability"), "chart-point-posterior"),
        ):
            if value is not None:
                point_markers.append(f'<circle cx="{x:.1f}" cy="{y_at(value):.1f}" r="5" class="{class_name}" />')
        grid_lines.append(
            f'<text x="{x:.1f}" y="{height - 16:.1f}" text-anchor="middle" class="chart-axis-label">R{point["round"]}</text>'
        )

    return f"""
      <svg class="trajectory-chart" viewBox="0 0 {width} {height}" role="img" aria-label="Belief trajectory chart">
        <rect x="0" y="0" width="{width}" height="{height}" fill="transparent"></rect>
        {''.join(grid_lines)}
        {''.join(readiness_bars)}
        <polyline points="{polyline(prior_values)}" class="chart-line-prior" />
        <polyline points="{polyline(posterior_values)}" class="chart-line-posterior" />
        {''.join(point_markers)}
      </svg>
    """


def confidence_phrase_zh(code: str | None) -> str:
    return {
        "low": "把握偏低",
        "medium": "把握一般",
        "medium-high": "把握较高",
    }.get(code, code or "把握一般")


def confidence_phrase_en(code: str | None) -> str:
    return {
        "low": "lower confidence",
        "medium": "moderate confidence",
        "medium-high": "stronger confidence",
    }.get(code, code or "moderate confidence")


def stability_phrase_zh(code: str | None) -> str:
    return {
        "stable": "换一组合理假设，结论大体不变",
        "mixed": "换一组合理假设，结论可能有轻微摆动",
        "unstable": "换一组合理假设，结论可能明显变化",
    }.get(code, code or "结论稳定性一般")


def stability_phrase_en(code: str | None) -> str:
    return {
        "stable": "The conclusion mostly holds under reasonable alternative assumptions.",
        "mixed": "The conclusion shifts a bit under reasonable alternative assumptions.",
        "unstable": "The conclusion can change materially under reasonable alternative assumptions.",
    }.get(code, code or "The conclusion has moderate stability.")


def classify_decision_state(report: dict) -> str:
    recommendation = report["decision"]["recommended_action"]
    if any(token in recommendation for token in ("暂停", "暂缓", "放弃", "延后", "改到")):
        return "hold"
    if any(token in recommendation for token in ("先", "确认", "测试", "试点", "验证", "再订")):
        return "confirm"
    return "proceed"


def evidence_ranked_pairs(request: dict, report: dict) -> list[tuple[float, dict, dict]]:
    pairs = []
    for zh_item, source_item in zip(report["evidence"], request.get("evidence", [])):
        lr = max(float(zh_item.get("likelihood_ratio") or 1.0), 1e-6)
        discount = max(float(zh_item.get("dependency_discount") or 1.0), 0.0)
        strength = abs(math.log(lr)) * discount
        pairs.append((strength, zh_item, source_item))
    return sorted(pairs, key=lambda item: item[0], reverse=True)


def top_reason_cards(request: dict, report: dict, direction_code: str, limit: int = 2) -> list[dict[str, str]]:
    cards: list[dict[str, str]] = []
    for _, zh_item, source_item in evidence_ranked_pairs(request, report):
        if zh_item.get("direction_code") != direction_code:
            continue
        cards.append(
            {
                "title_zh": zh_item["name"],
                "title_en": request_text(source_item.get("name"), "en"),
                "body_zh": zh_item.get("summary") or zh_item["name"],
                "body_en": request_text(source_item.get("summary"), "en") or request_text(source_item.get("name"), "en"),
            }
        )
        if len(cards) >= limit:
            break
    return cards


def alternative_cards(request: dict, report: dict) -> list[dict[str, str]]:
    ranking = report["decision"]["expected_value_ranking"]
    if len(ranking) <= 1:
        return []
    top_value = float(ranking[0]["expected_value"])
    cards = []
    for item in ranking[1:3]:
        gap = top_value - float(item["expected_value"])
        cards.append(
            {
                "name_zh": item["name"],
                "name_en": recommended_action_en(request, item["name"]),
                "reason_zh": f"它目前的综合收益比首选方案低 {fmt_num(gap)}，所以暂时排在后面。",
                "reason_en": f"It currently trails the recommended option by {fmt_num(gap)} in expected value, so it stays behind for now.",
            }
        )
    return cards


def plain_language_pack(request: dict, report: dict) -> dict[str, Any]:
    state = classify_decision_state(report)
    state_copy = DECISION_STATES[state]
    confidence_code = report["summary"].get("confidence_code")
    stability_code = report["sensitivity"].get("conclusion_stability_code")
    process = report.get("conversation_process") or {}
    deadline = report["question"].get("decision_deadline")
    success_metric = report["question"].get("success_metric") or "-"
    time_horizon = report["question"].get("time_horizon") or "-"
    next_info_zh = report["next_information"]["recommended_next_information"]
    next_info_en = first_next_information(request, "en")
    recommendation_zh = report["decision"]["recommended_action"]
    recommendation_en = recommended_action_en(request, recommendation_zh)

    if state == "proceed":
        step3_zh = (
            f"如果到 {deadline} 前没有出现新的明显负面信息，就按这个方案执行。"
            if deadline and deadline != "-"
            else "如果关键条件都确认无误，就按这个方案执行。"
        )
        step3_en = (
            f"If no major new negative signal shows up before {deadline}, execute this plan."
            if deadline and deadline != "-"
            else "If the key conditions check out, go ahead and execute this plan."
        )
    elif state == "confirm":
        step3_zh = (
            f"如果到 {deadline} 前关键条件确认无误，就推进；如果确认不了，就先不要做更重的动作。"
            if deadline and deadline != "-"
            else "如果关键条件确认无误，就推进；如果确认不了，就先不要做更重的动作。"
        )
        step3_en = (
            f"If the key conditions are confirmed before {deadline}, move ahead; if not, do not escalate to a heavier commitment yet."
            if deadline and deadline != "-"
            else "If the key conditions are confirmed, move ahead; if not, do not escalate to a heavier commitment yet."
        )
    else:
        step3_zh = "除非拿到新的强证据，否则先不要投入更多时间或预算。"
        step3_en = "Do not commit more time or budget unless stronger evidence appears."

    summary_zh = (
        f"{state_copy['tone_zh']} 目前的建议属于{confidence_phrase_zh(confidence_code)}，而且{stability_phrase_zh(stability_code)}。"
    )
    summary_en = (
        f"{state_copy['tone_en']} The current recommendation carries {confidence_phrase_en(confidence_code)}, and {stability_phrase_en(stability_code)}"
    )
    if process:
        gate_zh = (
            f"多轮对话后的决策准备度约为 {fmt_pct(process.get('final_readiness'))}，"
            f"{'已经' if process.get('decision_ready') else '还没有'}达到可决策状态。"
        )
        gate_en = (
            f"After the multi-turn loop, decision readiness is about {fmt_pct(process.get('final_readiness'))}, "
            f"and the case has {'reached' if process.get('decision_ready') else 'not yet reached'} a ready-to-decide state."
        )
    else:
        gate_zh = "这份判断基于当前一次性输入，没有额外的多轮过程记录。"
        gate_en = "This judgment is based on the current single-pass input and does not include a multi-turn process log."

    return {
        "state": state,
        "label_zh": state_copy["label_zh"],
        "label_en": state_copy["label_en"],
        "summary_zh": summary_zh,
        "summary_en": summary_en,
        "recommendation_zh": recommendation_zh,
        "recommendation_en": recommendation_en,
        "action_steps_zh": [
            recommendation_zh,
            next_info_zh,
            step3_zh,
        ],
        "action_steps_en": [
            recommendation_en,
            next_info_en,
            step3_en,
        ],
        "premises_zh": [
            f"时间窗口：{time_horizon}",
            f"成立条件：{success_metric}",
            f"最晚决策时间：{deadline}" if deadline and deadline != "-" else "最晚决策时间：尽快确认关键条件",
        ],
        "premises_en": [
            f"Time window: {request_text(request.get('time_horizon'), 'en')}",
            f"Condition for acting: {request_text(request.get('success_metric'), 'en')}",
            f"Decision deadline: {deadline}" if deadline and deadline != "-" else "Decision deadline: confirm the key conditions soon",
        ],
        "support_cards": top_reason_cards(request, report, "support"),
        "risk_cards": top_reason_cards(request, report, "against"),
        "alternatives": alternative_cards(request, report),
        "decision_gate_zh": gate_zh,
        "decision_gate_en": gate_en,
    }


def build_markdown(request: dict, report: dict) -> str:
    plain = plain_language_pack(request, report)
    process = report.get("conversation_process") or {}
    lines = [
        f"# 贝叶斯决策报告：{report['title']}",
        "",
        "> 本报告由 `yao-bayesian-skill` 基于同一结构化输入自动生成；JSON、Markdown、HTML、PDF、Word 版本保持同步。",
        "",
        "## 1. 先说结论",
        f"- 结论标签：{plain['label_zh']}",
        f"- 一句话判断：{report['summary']['one_sentence']}",
        f"- 人话解释：{plain['summary_zh']}",
        "",
        "## 2. 你现在该怎么做",
        f"1. {plain['action_steps_zh'][0]}",
        f"2. {plain['action_steps_zh'][1]}",
        f"3. {plain['action_steps_zh'][2]}",
        "",
        "## 3. 这份建议成立的前提",
        *[f"- {item}" for item in plain["premises_zh"]],
        "",
        "## 4. 为什么不是另外两个选项",
    ]
    if plain["alternatives"]:
        lines.extend([f"- {item['name_zh']}：{item['reason_zh']}" for item in plain["alternatives"]])
    else:
        lines.append("- 当前没有足够的备选行动用于比较。")

    if process:
        lines.extend(
            [
                "",
                "## 5. 多轮对话过程与决策准备度",
                f"- 初始现状：{process.get('initial_state') or '-'}",
                f"- 对话轮次：{process.get('round_count')}",
                f"- 当前状态：{process.get('status')}",
                f"- 决策准备度：{fmt_pct(process.get('final_readiness'))}",
                f"- 决策阈值：{fmt_pct(process.get('decision_ready_threshold'))}",
                f"- 是否可以进入正式决策：{'可以' if process.get('decision_ready') else '还不可以'}",
                f"- 过程分析：{process.get('analysis')}",
                "",
                "| 轮次 | 阶段 | 起点概率 | 更新后概率 | 概率变化 | 决策准备度 | 中间判断 |",
                "|---:|---|---:|---:|---:|---:|---|",
            ]
        )
        for item in process.get("rounds", []):
            lines.append(
                f"| {item['round']} | {item['stage']} | {fmt_pct(item.get('prior_probability_before'))} | {fmt_pct(item.get('posterior_probability_after'))} | {fmt_pct(item.get('delta_probability'))} | {fmt_pct(item.get('decision_readiness'))} | {item.get('interim_judgment') or '-'} |"
            )
        if process.get("open_questions"):
            lines.extend(
                [
                    "",
                    "- 当前仍需关注的问题：",
                    *[f"  - {item}" for item in process["open_questions"]],
                ]
            )
        lines.extend(
            [
                "",
                "- 贝叶斯过程日志：",
            ]
        )
        for item in process.get("rounds", []):
            formula = (item.get("bayes_update") or {}).get("formula_zh")
            lines.append(f"  - 第 {item['round']} 轮：{formula or '本轮只记录了结果，没有完整公式参数。'}")

    lines.extend(
        [
            "",
            "## 6. 决策问题",
            f"- 决策问题：{report['question']['decision_question']}",
            f"- 要判断的假设：{report['question']['hypothesis']}",
            f"- 时间范围：{report['question']['time_horizon']}",
            f"- 决策截止时间：{report['question']['decision_deadline']}",
            f"- 成功标准：{report['question']['success_metric']}",
            f"- 领域：{report['question']['domain']}",
            "",
            "## 7. 先验设置",
            f"- 先验概率：{fmt_pct(report['prior']['probability'])}",
            f"- 先验分布：{report['prior']['distribution']}",
            f"- 先验区间：{fmt_interval(report['prior']['credible_interval'])}",
            f"- 先验来源等级：{report['prior']['source_quality']}",
            f"- 先验信心：{report['prior']['confidence']}",
            f"- 等效样本量：{fmt_num(report['prior']['equivalent_sample_size'])}",
            "",
        ]
    )
    if report["prior"]["source_summary"]:
        lines.append("- 先验来源：")
        lines.extend([f"  - {item}" for item in report["prior"]["source_summary"]])
        lines.append("")

    lines.extend(
        [
            "## 8. 证据摘要",
            "| 证据 | 可信度 | 方向 | 似然比 | 依赖折扣 | 摘要 |",
            "|---|---|---|---:|---:|---|",
        ]
    )
    for item in report["evidence"]:
        lines.append(
            f"| {item['name']} | {item['quality']} | {item['direction']} | {fmt_num(item['likelihood_ratio'])} | {fmt_num(item['dependency_discount'])} | {item['summary'] or '-'} |"
        )

    lines.extend(
        [
            "",
            "## 9. 贝叶斯更新",
            f"- 更新方法：{report['posterior']['method']}",
            f"- 后验概率：{fmt_pct(report['posterior']['probability'])}",
            f"- 后验区间：{fmt_interval(report['posterior']['credible_interval'])}",
            f"- 自然频率解释：{report['natural_frequency']['explanation']}",
            "",
            "## 10. 行动比较",
            f"- 推荐行动：{report['decision']['recommended_action']}",
            f"- 推荐理由：{report['decision']['reason']}",
            f"- 决策状态：{report['decision'].get('decision_status') or '未单独评估'}",
            f"- 决策准备度：{fmt_pct(report['decision'].get('decision_readiness'))}",
            "",
            "| 行动 | H 为真时效用 | H 为假时效用 | 期望值 | 行动阈值 |",
            "|---|---:|---:|---:|---:|",
        ]
    )
    for action in report["decision"]["expected_value_ranking"]:
        lines.append(
            f"| {action['name']} | {fmt_num(action['utility_if_h_true'])} | {fmt_num(action['utility_if_h_false'])} | {fmt_num(action['expected_value'])} | {fmt_pct(action['action_threshold'])} |"
        )

    lines.extend(
        [
            "",
            "## 11. 敏感性分析",
            f"- 后验范围：{fmt_interval(report['sensitivity']['posterior_range'])}",
            f"- 结论稳定性：{report['sensitivity']['conclusion_stability']}",
            "",
            "| 先验概率 | LR 强度系数 | 后验概率 | 推荐行动 |",
            "|---:|---:|---:|---|",
        ]
    )
    for scenario in report["sensitivity"]["scenarios"]:
        lines.append(
            f"| {fmt_pct(scenario.get('prior_probability'))} | {fmt_num(scenario.get('lr_power'))} | {fmt_pct(scenario.get('posterior_probability'))} | {scenario.get('recommended_action') or '-'} |"
        )

    lines.extend(
        [
            "",
            "## 12. 下一步最有价值的信息",
            f"- 推荐下一步信息：{report['next_information']['recommended_next_information']}",
            f"- 判断原因：{report['next_information']['reason']}",
            "",
            "## 13. 风险与注意事项",
        ]
    )
    if report["warnings"]:
        lines.extend([f"- {warning}" for warning in report["warnings"]])
    else:
        lines.append("- 当前没有额外警示项。")

    lines.extend(
        [
            "",
            "## 14. Skill 流程",
        ]
    )
    for index, item in enumerate(SKILL_FLOW, start=1):
        lines.append(f"- {index}. {item['zh']}")

    lines.extend(
        [
            "",
            "## 15. Skill 能力",
        ]
    )
    for capability in SKILL_CAPABILITIES:
        lines.append(f"- {capability['zh_title']}：{capability['zh_body']}")

    lines.extend(
        [
            "",
            "## 16. 自动生成说明",
            "- 本报告不是手写示例，而是由同一个结构化输入自动渲染出来的正式输出。",
            "- HTML 提供中英双语一键切换，PDF 和 Word 默认采用简体中文主报告。",
            f"- {plain['decision_gate_zh']}",
        ]
    )
    return "\n".join(lines).strip() + "\n"


def build_html(request: dict, report: dict) -> str:
    css = CSS_PATH.read_text(encoding="utf-8")
    plain = plain_language_pack(request, report)
    process = report.get("conversation_process") or {}
    raw_rounds = conversation_raw_rounds(request)
    title_zh = report["title"]
    title_en = request_text(request.get("title"), "en")
    decision_zh = report["question"]["decision_question"]
    decision_en = request_text(request.get("question"), "en")
    hypothesis_zh = report["question"]["hypothesis"]
    hypothesis_en = request_text(request.get("hypothesis"), "en")
    horizon_zh = report["question"]["time_horizon"]
    horizon_en = request_text(request.get("time_horizon"), "en")
    success_zh = report["question"]["success_metric"]
    success_en = request_text(request.get("success_metric"), "en")
    domain_zh = report["question"]["domain"]
    domain_en = request_text(request.get("domain"), "en")
    prior_sources_zh = report["prior"]["source_summary"]
    prior_sources_en = localize_items(request.get("prior", {}).get("source_summary"), "en")
    recommended_zh = report["decision"]["recommended_action"]
    recommended_en = recommended_action_en(request, recommended_zh)
    ranking = report["decision"]["expected_value_ranking"]
    top_threshold = ranking[0].get("action_threshold") if ranking else None
    support_cards_html = "".join(
        """
        <article class="reason-card">
          <div class="reason-label">{label}</div>
          <h3>{title}</h3>
          <p>{body}</p>
        </article>
        """.format(
            label=dual_html("支持判断", "Supports the call"),
            title=dual_html(item["title_zh"], item["title_en"]),
            body=dual_html(item["body_zh"], item["body_en"]),
        )
        for item in plain["support_cards"]
    ) or """
        <article class="reason-card">
          <div class="reason-label"><span class="lang-zh">支持判断</span><span class="lang-en">Supports the call</span></div>
          <p><span class="lang-zh">当前没有足够的正向证据卡片。</span><span class="lang-en">There are not enough positive evidence cards at the moment.</span></p>
        </article>
    """
    risk_cards_html = "".join(
        """
        <article class="reason-card risk">
          <div class="reason-label">{label}</div>
          <h3>{title}</h3>
          <p>{body}</p>
        </article>
        """.format(
            label=dual_html("拉低判断", "Pulls the call down"),
            title=dual_html(item["title_zh"], item["title_en"]),
            body=dual_html(item["body_zh"], item["body_en"]),
        )
        for item in plain["risk_cards"]
    ) or """
        <article class="reason-card risk">
          <div class="reason-label"><span class="lang-zh">拉低判断</span><span class="lang-en">Pulls the call down</span></div>
          <p><span class="lang-zh">当前没有足够的负向证据卡片。</span><span class="lang-en">There are not enough negative evidence cards at the moment.</span></p>
        </article>
    """
    alternative_cards_html = "".join(
        """
        <li>
          <strong>{name}</strong>
          <span>{reason}</span>
        </li>
        """.format(
            name=dual_html(item["name_zh"], item["name_en"]),
            reason=dual_html(item["reason_zh"], item["reason_en"]),
        )
        for item in plain["alternatives"]
    ) or """
        <li><span class="lang-zh">当前没有足够的备选动作可比较。</span><span class="lang-en">There are not enough alternate actions to compare right now.</span></li>
    """
    action_steps_html = "".join(
        f"<li>{dual_html(plain['action_steps_zh'][index], plain['action_steps_en'][index])}</li>"
        for index in range(len(plain["action_steps_zh"]))
    )
    premise_items_html = "".join(
        f"<li>{dual_html(plain['premises_zh'][index], plain['premises_en'][index])}</li>"
        for index in range(len(plain["premises_zh"]))
    )
    process_badge_html = ""
    if process:
        process_badge_html = f"""
          <span class="pill">{dual_html(process.get("status") or "-", decision_status_en(report))}</span>
          <span class="pill">{dual_html(f"对话 {process.get('round_count')} 轮", f"{process.get('round_count')} rounds")}</span>
        """
    evidence_rows = []
    for zh_item, source_item in zip(report["evidence"], request.get("evidence", [])):
        evidence_rows.append(
            """
            <tr>
              <td>{name}</td>
              <td><span class="pill">{quality}</span></td>
              <td>{direction}</td>
              <td>{lr}</td>
              <td>{discount}</td>
              <td>{summary}</td>
              <td>{notes}</td>
            </tr>
            """.format(
                name=dual_html(zh_item["name"], request_text(source_item.get("name"), "en")),
                quality=html_text(zh_item["quality"]),
                direction=dual_html(zh_item["direction"], direction_en(zh_item)),
                lr=html_text(fmt_num(zh_item["likelihood_ratio"])),
                discount=html_text(fmt_num(zh_item["dependency_discount"])),
                summary=dual_html(zh_item.get("summary") or "-", request_text(source_item.get("summary"), "en")),
                notes=dual_html(zh_item.get("notes") or "-", request_text(source_item.get("notes"), "en")),
            )
        )

    action_rows = []
    for item in ranking:
        action_rows.append(
            """
            <tr>
              <td>{name}</td>
              <td>{utility_true}</td>
              <td>{utility_false}</td>
              <td>{expected_value}</td>
              <td>{threshold}</td>
            </tr>
            """.format(
                name=dual_html(item["name"], recommended_action_en(request, item["name"])),
                utility_true=html_text(fmt_num(item["utility_if_h_true"])),
                utility_false=html_text(fmt_num(item["utility_if_h_false"])),
                expected_value=html_text(fmt_num(item["expected_value"])),
                threshold=html_text(fmt_pct(item["action_threshold"])),
            )
        )

    sensitivity_rows = []
    for item in report["sensitivity"]["scenarios"]:
        sensitivity_rows.append(
            """
            <tr>
              <td>{prior}</td>
              <td>{lr_power}</td>
              <td>{posterior}</td>
              <td>{action}</td>
            </tr>
            """.format(
                prior=html_text(fmt_pct(item.get("prior_probability"))),
                lr_power=html_text(fmt_num(item.get("lr_power"))),
                posterior=html_text(fmt_pct(item.get("posterior_probability"))),
                action=dual_html(item.get("recommended_action") or "-", recommended_action_en(request, item.get("recommended_action") or "-")),
            )
        )

    prior_source_items = "".join(
        f"<li>{dual_html(prior_sources_zh[index], prior_sources_en[index] if index < len(prior_sources_en) else prior_sources_zh[index])}</li>"
        for index in range(len(prior_sources_zh))
    )
    warning_items = "".join(f"<li>{dual_html(item, warning_en(item))}</li>" for item in report["warnings"]) or f"<li>{dual_html('当前没有额外警示项。', 'There are no additional warnings at the moment.')}</li>"
    workflow_cards = "".join(
        """
        <div class="workflow-step">
          <div class="step-index">{index}</div>
          <div class="step-text">{label}</div>
        </div>
        """.format(index=html_text(f"Step {index}"), label=dual_html(item["zh"], item["en"]))
        for index, item in enumerate(SKILL_FLOW, start=1)
    )
    capability_cards = "".join(
        """
        <div class="capability-card">
          <h3>{title}</h3>
          <p>{body}</p>
        </div>
        """.format(title=dual_html(item["zh_title"], item["en_title"]), body=dual_html(item["zh_body"], item["en_body"]))
        for item in SKILL_CAPABILITIES
    )
    conversation_rounds_html = ""
    conversation_section_html = ""
    if process:
        for index, item in enumerate(process.get("rounds", [])):
            raw_round = raw_rounds[index] if index < len(raw_rounds) else {}
            formula = item.get("bayes_update") or {}
            conversation_rounds_html += """
            <article class="round-card">
              <div class="round-card-top">
                <div>
                  <div class="section-kicker">{round_label}</div>
                  <h3>{stage}</h3>
                </div>
                <div class="round-badge">{readiness}</div>
              </div>
              <div class="round-metrics">
                <div class="metric-chip"><strong>{prior_label}</strong><span>{prior_value}</span></div>
                <div class="metric-chip"><strong>{posterior_label}</strong><span>{posterior_value}</span></div>
                <div class="metric-chip"><strong>{delta_label}</strong><span>{delta_value}</span></div>
                <div class="metric-chip"><strong>{ready_label}</strong><span>{ready_value}</span></div>
              </div>
              <p><strong>{user_label}</strong> {user_value}</p>
              <p><strong>{focus_label}</strong> {focus_value}</p>
              <p><strong>{judgment_label}</strong> {judgment_value}</p>
              <div class="callout process-callout">
                <p><strong>{formula_label}</strong> {formula_value}</p>
              </div>
              <div class="round-two-col">
                <div>
                  <h4>{new_info_label}</h4>
                  <ul>{new_info_items}</ul>
                </div>
                <div>
                  <h4>{next_q_label}</h4>
                  <ul>{next_q_items}</ul>
                </div>
              </div>
            </article>
            """.format(
                round_label=dual_html(f"第 {item['round']} 轮", f"Round {item['round']}"),
                stage=dual_html(item.get("stage") or "-", conversation_round_en(raw_round, "stage")),
                readiness=dual_html(item.get("decision_readiness_label") or "-", readiness_en(item.get("decision_readiness_label") or "-")),
                prior_label=dual_html("起点概率", "Starting belief"),
                prior_value=html_text(fmt_pct(item.get("prior_probability_before"))),
                posterior_label=dual_html("更新后概率", "Updated belief"),
                posterior_value=html_text(fmt_pct(item.get("posterior_probability_after"))),
                delta_label=dual_html("概率变化", "Change"),
                delta_value=html_text(fmt_pct(item.get("delta_probability"))),
                ready_label=dual_html("准备度", "Readiness"),
                ready_value=html_text(fmt_pct(item.get("decision_readiness"))),
                user_label=dual_html("用户补充了什么：", "What the user added:"),
                user_value=dual_html(item.get("user_input_summary") or "-", conversation_round_en(raw_round, "user_input_summary")),
                focus_label=dual_html("这一轮我在判断什么：", "What the skill focused on:"),
                focus_value=dual_html(item.get("assistant_focus") or "-", conversation_round_en(raw_round, "assistant_focus")),
                judgment_label=dual_html("中间判断：", "Interim judgment:"),
                judgment_value=dual_html(item.get("interim_judgment") or "-", conversation_round_en(raw_round, "interim_judgment")),
                formula_label=dual_html("贝叶斯变化：", "Bayesian change:"),
                formula_value=dual_html(formula.get("formula_zh") or "-", formula.get("formula_en") or "-"),
                new_info_label=dual_html("新增信息", "New information"),
                new_info_items=conversation_dual_list_items(item, raw_round, "new_information")
                or f"<li>{dual_html('本轮没有新增结构化信息。', 'No new structured information was recorded for this round.')}</li>",
                next_q_label=dual_html("下一步追问 / 剩余缺口", "Next questions / remaining gaps"),
                next_q_items=conversation_dual_list_items(item, raw_round, "assistant_next_questions", fallback_key="missing_information")
                or f"<li>{dual_html('本轮后已经没有剩余关键缺口。', 'There are no remaining critical gaps after this round.')}</li>",
            )

        conversation_section_html = f"""
          <div class="grid compact-grid">
            <div class="metric-card">
              <div class="metric-label">{dual_html("对话轮次", "Rounds")}</div>
              <div class="metric-value">{html_text(fmt_num(process.get("round_count")))}</div>
            </div>
            <div class="metric-card">
              <div class="metric-label">{dual_html("决策准备度", "Decision readiness")}</div>
              <div class="metric-value">{html_text(fmt_pct(process.get("final_readiness")))}</div>
            </div>
            <div class="metric-card">
              <div class="metric-label">{dual_html("决策阈值", "Readiness threshold")}</div>
              <div class="metric-value">{html_text(fmt_pct(process.get("decision_ready_threshold")))}</div>
            </div>
            <div class="metric-card">
              <div class="metric-label">{dual_html("当前状态", "Current status")}</div>
              <div class="metric-value metric-text">{dual_html(process.get("status") or "-", decision_status_en(report))}</div>
            </div>
          </div>
          <p>{dual_html(process.get("analysis") or "-", process.get("analysis_en") or "-")}</p>
          <p class="muted-note">{dual_html(plain['decision_gate_zh'], plain['decision_gate_en'])}</p>
          <div class="trajectory-shell">
            <div class="trajectory-head">
              <h3>{dual_html("变化图表", "Trajectory chart")}</h3>
              <div class="trajectory-legend">
                <span class="legend-item"><span class="legend-dot legend-dot-prior"></span>{dual_html("起点概率", "Prior")}</span>
                <span class="legend-item"><span class="legend-dot legend-dot-posterior"></span>{dual_html("更新后概率", "Posterior")}</span>
                <span class="legend-item"><span class="legend-bar"></span>{dual_html("决策准备度", "Readiness")}</span>
              </div>
            </div>
            {conversation_chart_svg(report)}
          </div>
          <div class="round-grid">{conversation_rounds_html}</div>
        """
    nav_links = "".join(
        f'<a class="menu-link{" nav-pro-only" if anchor not in {"summary", "action-plan"} else ""}" href="#{anchor}">{dual_html(zh, en)}</a>'
        for anchor, zh, en in TOP_NAV
    )

    prior_section_html = f"""
      <p><strong>{dual_html("来源等级：", "Source quality:")}</strong> {html_text(report["prior"]["source_quality"])}</p>
      <p><strong>{dual_html("先验区间：", "Prior interval:")}</strong> {html_text(fmt_interval(report["prior"]["credible_interval"]))}</p>
      <ul>{prior_source_items}</ul>
    """

    evidence_section_html = f"""
      <div class="table-wrap">
        <table>
          <thead>
            <tr>
              <th>{dual_html("证据", "Evidence")}</th>
              <th>{dual_html("等级", "Grade")}</th>
              <th>{dual_html("方向", "Direction")}</th>
              <th>{dual_html("似然比", "LR")}</th>
              <th>{dual_html("依赖折扣", "Dependency")}</th>
              <th>{dual_html("摘要", "Summary")}</th>
              <th>{dual_html("说明", "Notes")}</th>
            </tr>
          </thead>
          <tbody>
            {''.join(evidence_rows)}
          </tbody>
        </table>
      </div>
    """

    update_section_html = f"""
      <p><strong>{dual_html("更新方法：", "Method:")}</strong> {html_text(report["posterior"]["method"])}</p>
      <p><strong>{dual_html("后验概率：", "Posterior probability:")}</strong> {html_text(fmt_pct(report["posterior"]["probability"]))}</p>
      <p><strong>{dual_html("后验区间：", "Posterior interval:")}</strong> {html_text(fmt_interval(report["posterior"]["credible_interval"]))}</p>
      <div class="callout">
        <p>{dual_html(report["natural_frequency"]["explanation"], natural_frequency_en(report))}</p>
      </div>
    """

    sensitivity_section_html = f"""
      <p><strong>{dual_html("后验范围：", "Posterior range:")}</strong> {html_text(fmt_interval(report["sensitivity"]["posterior_range"]))}</p>
      <p><strong>{dual_html("结论稳定性：", "Conclusion stability:")}</strong> {dual_html(report["sensitivity"]["conclusion_stability"], stability_en(report))}</p>
      <div class="table-wrap">
        <table>
          <thead>
            <tr>
              <th>{dual_html("先验概率", "Prior")}</th>
              <th>{dual_html("LR 强度系数", "LR power")}</th>
              <th>{dual_html("后验概率", "Posterior")}</th>
              <th>{dual_html("推荐行动", "Recommended action")}</th>
            </tr>
          </thead>
          <tbody>
            {''.join(sensitivity_rows)}
          </tbody>
        </table>
      </div>
    """

    appendix_section_html = f"""
      <div class="appendix-block">
        <h3>{dual_html("Skill 流程", "Skill workflow")}</h3>
        <div class="workflow">{workflow_cards}</div>
      </div>
      <div class="appendix-block">
        <h3>{dual_html("Skill 能力", "Skill capabilities")}</h3>
        <div class="capability-grid">{capability_cards}</div>
      </div>
    """

    html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
  <meta charset="utf-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1" />
  <title>{html_text(title_zh)} | Yao Bayesian Skill</title>
  <style>
{css}
  </style>
</head>
<body data-lang="zh" data-view="simple">
  <header class="topbar">
    <div class="topbar-inner">
      <div class="brand-block">
        <div class="brand-mark">YAO</div>
        <div>
          <div class="brand-title">{dual_html("贝叶斯决策报告", "Bayesian Decision Report")}</div>
          <div class="brand-subtitle">{dual_html("自动生成 HTML / PDF / Word bundle", "Auto-generated HTML / PDF / Word bundle")}</div>
        </div>
      </div>
      <nav class="menu">{nav_links}</nav>
      <div class="view-toggle" aria-label="View mode">
        <button class="view-button is-active" type="button" data-view-button="simple">简版</button>
        <button class="view-button" type="button" data-view-button="pro">专业版</button>
      </div>
      <div class="lang-toggle" aria-label="Language switcher">
        <button class="lang-button is-active" type="button" data-lang-button="zh">中文</button>
        <button class="lang-button" type="button" data-lang-button="en">EN</button>
      </div>
    </div>
  </header>

  <main class="page">
    <section class="hero section-anchor-offset" id="summary">
      <div class="eyebrow">{dual_html("Yao Bayesian Skill", "Yao Bayesian Skill")}</div>
      {dual_tag("h1", title_zh, title_en)}
      <p class="summary-line">{dual_html(report["summary"]["one_sentence"], summary_en(request, report))}</p>
      <div class="decision-banner decision-{plain['state']}">
        <div class="decision-banner-label">{dual_html("给普通用户的结论", "Plain-language decision")}</div>
        <div class="decision-banner-state">{dual_html(plain["label_zh"], plain["label_en"])}</div>
        <p class="decision-banner-text">{dual_html(plain["summary_zh"], plain["summary_en"])}</p>
      </div>
      <div class="hero-meta">
        <span class="pill">{dual_html("自动生成", "Auto-generated")}</span>
        <span class="pill">{dual_html("中文主报告 + 英文切换", "Chinese-primary + English toggle")}</span>
        <span class="pill">{dual_html("默认简版只看结论和行动", "Simple mode shows only the conclusion and action")}</span>
        {process_badge_html}
      </div>
      <div class="grid hero-grid">
        <div class="metric-card">
          <div class="metric-label">{dual_html("后验概率", "Posterior")}</div>
          <div class="metric-value">{html_text(fmt_pct(report["summary"]["posterior_probability"]))}</div>
        </div>
        <div class="metric-card">
          <div class="metric-label">{dual_html("后验区间", "Interval")}</div>
          <div class="metric-value metric-text">{html_text(fmt_interval(report["summary"]["credible_interval"]))}</div>
        </div>
        <div class="metric-card">
          <div class="metric-label">{dual_html("推荐行动", "Recommended action")}</div>
          <div class="metric-value metric-text">{dual_html(recommended_zh, recommended_en)}</div>
        </div>
        <div class="metric-card">
          <div class="metric-label">{dual_html("建议信心", "Decision confidence")}</div>
          <div class="metric-value metric-text">{dual_html(report["summary"]["confidence"], CONFIDENCE_EN.get(report["summary"]["confidence_code"], report["summary"]["confidence_code"]))}</div>
        </div>
        <div class="metric-card">
          <div class="metric-label">{dual_html("首选行动阈值", "Top-action threshold")}</div>
          <div class="metric-value">{html_text(fmt_pct(top_threshold))}</div>
        </div>
        <div class="metric-card">
          <div class="metric-label">{dual_html("决策准备度", "Decision readiness")}</div>
          <div class="metric-value">{html_text(fmt_pct(report["summary"].get("decision_readiness")))}</div>
        </div>
        <div class="metric-card">
          <div class="metric-label">{dual_html("可决策状态", "Decision status")}</div>
          <div class="metric-value metric-text">{dual_html(report["summary"].get("decision_status") or "未单独评估", decision_status_en(report))}</div>
        </div>
      </div>
      <div class="meter" aria-hidden="true">
        <span style="width: {min(100.0, max(0.0, float(report['summary']['posterior_probability']) * 100.0)):.1f}%"></span>
      </div>
    </section>

    <section class="section executive-panel section-anchor-offset" id="action-plan">
      <div class="section-kicker">{dual_html("先做执行摘要", "Executive summary first")}</div>
      <div class="executive-grid">
        <article class="executive-card executive-primary">
          <h2>{dual_html("你现在该怎么做", "What you should do now")}</h2>
          <ol class="action-list">{action_steps_html}</ol>
        </article>
        <article class="executive-card">
          <h2>{dual_html("这份建议成立的前提", "What this recommendation depends on")}</h2>
          <ul class="premise-list">{premise_items_html}</ul>
          <p class="muted-note">{dual_html(plain["decision_gate_zh"], plain["decision_gate_en"])}</p>
        </article>
      </div>
    </section>

    <section class="section pro-only section-anchor-offset" id="why">
      <div class="section-kicker">{dual_html("先讲人话，再讲模型", "Human explanation before model detail")}</div>
      <h2>{dual_html("为什么这样判断", "Why this is the recommendation")}</h2>
      <div class="reason-grid">
        {support_cards_html}
        {risk_cards_html}
      </div>
      <div class="comparison-card">
        <h3>{dual_html("为什么不是另外两个选项", "Why not the other options")}</h3>
        <ul class="comparison-list">{alternative_cards_html}</ul>
      </div>
    </section>

    <section class="section pro-only section-anchor-offset" id="decision">
      {dual_tag("h2", "决策问题", "Decision framing")}
      <p><strong>{dual_html("决策问题：", "Decision question:")}</strong> {dual_html(decision_zh, decision_en)}</p>
      <p><strong>{dual_html("当前现状：", "Current state:")}</strong> {dual_html(request_text(request.get("current_state"), "zh"), request_text(request.get("current_state"), "en"))}</p>
      <p><strong>{dual_html("假设：", "Hypothesis:")}</strong> {dual_html(hypothesis_zh, hypothesis_en)}</p>
      <p><strong>{dual_html("时间范围：", "Time horizon:")}</strong> {dual_html(horizon_zh, horizon_en)}</p>
      <p><strong>{dual_html("成功标准：", "Success metric:")}</strong> {dual_html(success_zh, success_en)}</p>
      <p><strong>{dual_html("领域：", "Domain:")}</strong> {dual_html(domain_zh, domain_en)}</p>
      <p class="muted-note">{dual_html("这份报告由 skill 自动把问题、证据、行动和阈值对齐成统一决策格式。", "This report is auto-structured by the skill so the question, evidence, actions, and thresholds stay aligned.")}</p>
    </section>

    {fold_section("conversation", "多轮对话过程", "Conversation process", conversation_section_html or f"<p>{dual_html('当前没有记录多轮过程；报告只展示最终结果。', 'There is no multi-turn process log for this report, so only the final result is shown.')}</p>", "默认折叠", "Collapsed by default")}

    {fold_section(
        "prior",
        "先验设置",
        "Prior setup",
        f'''
          <div class="grid compact-grid process-grid">
        <div class="metric-card">
          <div class="metric-label">{dual_html("先验概率", "Prior probability")}</div>
          <div class="metric-value">{html_text(fmt_pct(report["prior"]["probability"]))}</div>
        </div>
        <div class="metric-card">
          <div class="metric-label">{dual_html("先验分布", "Prior distribution")}</div>
          <div class="metric-value metric-text">{html_text(report["prior"]["distribution"])}</div>
        </div>
        <div class="metric-card">
          <div class="metric-label">{dual_html("等效样本量", "Equivalent sample size")}</div>
          <div class="metric-value">{html_text(fmt_num(report["prior"]["equivalent_sample_size"]))}</div>
        </div>
        <div class="metric-card">
          <div class="metric-label">{dual_html("先验信心", "Prior confidence")}</div>
          <div class="metric-value metric-text">{dual_html(report["prior"]["confidence"], prior_confidence_en(request, report))}</div>
        </div>
      </div>
      {prior_section_html}
      '''
    )}

    {fold_section("evidence", "证据摘要", "Evidence summary", evidence_section_html)}

    {fold_section("update", "概率判断", "Probability view", update_section_html)}

    <section class="section pro-only section-anchor-offset" id="actions">
      {dual_tag("h2", "行动比较", "Action comparison")}
      <p><strong>{dual_html("推荐行动：", "Recommended action:")}</strong> {dual_html(recommended_zh, recommended_en)}</p>
      <p>{dual_html(report["decision"]["reason"], decision_reason_en(report))}</p>
      <div class="table-wrap">
        <table>
          <thead>
            <tr>
              <th>{dual_html("行动", "Action")}</th>
              <th>{dual_html("H 为真时效用", "Utility if H is true")}</th>
              <th>{dual_html("H 为假时效用", "Utility if H is false")}</th>
              <th>{dual_html("期望值", "Expected value")}</th>
              <th>{dual_html("行动阈值", "Action threshold")}</th>
            </tr>
          </thead>
          <tbody>
            {''.join(action_rows)}
          </tbody>
        </table>
      </div>
      <h3>{dual_html("下一步最有价值的信息", "Most valuable next information")}</h3>
      <p>{dual_html(report["next_information"]["recommended_next_information"], first_next_information(request, "en"))}</p>
      <p class="muted-note">{dual_html(report["next_information"]["reason"], next_information_reason_en(report))}</p>
    </section>

    {fold_section("sensitivity", "敏感性分析", "Sensitivity analysis", sensitivity_section_html)}

    <section class="section pro-only section-anchor-offset" id="warnings">
      {dual_tag("h2", "风险与注意事项", "Warnings and caveats")}
      <ul class="warning-list">{warning_items}</ul>
      <p class="footer-note">{dual_html("本 HTML 只是同一份自动化 bundle 的一个可视化视图；PDF 和 Word 与其共享同一个结构化计算结果。", "This HTML report is only one view of the same automated bundle; PDF and Word share the same structured calculation result.")}</p>
    </section>

    {fold_section("workflow", "附录", "Appendix", appendix_section_html)}
  </main>

  <script>
    (function () {{
      const root = document.body;
      const buttons = Array.from(document.querySelectorAll("[data-lang-button]"));
      const viewButtons = Array.from(document.querySelectorAll("[data-view-button]"));
      function setLang(lang) {{
        root.dataset.lang = lang;
        buttons.forEach((button) => {{
          button.classList.toggle("is-active", button.dataset.langButton === lang);
        }});
        try {{
          window.localStorage.setItem("yao-bayesian-report-lang", lang);
        }} catch (error) {{
        }}
      }}
      function setView(view) {{
        root.dataset.view = view;
        viewButtons.forEach((button) => {{
          button.classList.toggle("is-active", button.dataset.viewButton === view);
        }});
        try {{
          window.localStorage.setItem("yao-bayesian-report-view", view);
        }} catch (error) {{
        }}
      }}
      buttons.forEach((button) => {{
        button.addEventListener("click", () => setLang(button.dataset.langButton));
      }});
      viewButtons.forEach((button) => {{
        button.addEventListener("click", () => setView(button.dataset.viewButton));
      }});
      document.querySelectorAll('a.menu-link').forEach((link) => {{
        link.addEventListener('click', () => {{
          const href = link.getAttribute('href') || '';
          const target = href.startsWith('#') ? document.getElementById(href.slice(1)) : null;
          if (target && target.tagName.toLowerCase() === 'details') {{
            target.open = true;
          }}
        }});
      }});
      try {{
        const saved = window.localStorage.getItem("yao-bayesian-report-lang");
        if (saved === "en" || saved === "zh") {{
          setLang(saved);
        }}
      }} catch (error) {{
      }}
      try {{
        const savedView = window.localStorage.getItem("yao-bayesian-report-view");
        if (savedView === "pro" || savedView === "simple") {{
          setView(savedView);
        }} else {{
          setView("simple");
        }}
      }} catch (error) {{
        setView("simple");
      }}
      if (!root.dataset.lang) {{
        setLang("zh");
      }}
    }})();
  </script>
</body>
</html>
"""
    return html


def register_pdf_fonts() -> None:
    try:
        pdfmetrics.getFont("STSong-Light")
    except KeyError:
        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))


def pdf_styles():
    register_pdf_fonts()
    styles = getSampleStyleSheet()
    title = ParagraphStyle(
        "ReportTitle",
        parent=styles["Title"],
        fontName="STSong-Light",
        fontSize=22,
        leading=28,
        textColor=colors.HexColor("#141413"),
        spaceAfter=10,
    )
    heading = ParagraphStyle(
        "SectionHeading",
        parent=styles["Heading2"],
        fontName="STSong-Light",
        fontSize=15,
        leading=20,
        textColor=colors.HexColor("#1B365D"),
        spaceBefore=10,
        spaceAfter=6,
    )
    body = ParagraphStyle(
        "BodyCN",
        parent=styles["BodyText"],
        fontName="STSong-Light",
        fontSize=10.5,
        leading=16,
        textColor=colors.HexColor("#141413"),
        spaceAfter=5,
    )
    small = ParagraphStyle(
        "SmallCN",
        parent=body,
        fontSize=9,
        leading=13,
        textColor=colors.HexColor("#5e5d59"),
    )
    return title, heading, body, small


def paragraph(text: str, style: ParagraphStyle) -> Paragraph:
    return Paragraph(escape(text).replace("\n", "<br/>"), style)


def make_pdf_table(rows: list[list[str]], widths: list[float]) -> Table:
    table = Table(rows, colWidths=widths, repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("FONTNAME", (0, 0), (-1, -1), "STSong-Light"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("LEADING", (0, 0), (-1, -1), 12),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#EEF2F7")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#141413")),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#e8e5da")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    return table


def build_pdf(request: dict, report: dict, output_path: Path) -> None:
    plain = plain_language_pack(request, report)
    process = report.get("conversation_process") or {}
    title_style, heading_style, body_style, small_style = pdf_styles()
    story = []
    story.append(paragraph(f"贝叶斯决策报告：{report['title']}", title_style))
    story.append(paragraph("本报告由 yao-bayesian-skill 自动生成，PDF 与 HTML、Word、JSON 使用同一套结构化计算结果。", small_style))
    story.append(Spacer(1, 4 * mm))

    story.append(paragraph("先说结论", heading_style))
    story.append(paragraph(report["summary"]["one_sentence"], body_style))
    story.append(paragraph(plain["summary_zh"], body_style))

    story.append(paragraph("你现在该怎么做", heading_style))
    for index, step in enumerate(plain["action_steps_zh"], start=1):
        story.append(paragraph(f"{index}. {step}", body_style))

    story.append(paragraph("这份建议成立的前提", heading_style))
    for item in plain["premises_zh"]:
        story.append(paragraph(f"• {item}", body_style))
    story.append(paragraph(plain["decision_gate_zh"], small_style))

    story.append(paragraph("为什么不是另外两个选项", heading_style))
    if plain["alternatives"]:
        for item in plain["alternatives"]:
            story.append(paragraph(f"{item['name_zh']}：{item['reason_zh']}", body_style))
    else:
        story.append(paragraph("当前没有足够的备选行动用于比较。", body_style))

    if process:
        story.append(paragraph("多轮对话过程与决策准备度", heading_style))
        story.append(paragraph(process.get("analysis") or "-", body_style))
        process_lines = [
            f"当前状态：{process.get('status')}",
            f"决策准备度：{fmt_pct(process.get('final_readiness'))}",
            f"决策阈值：{fmt_pct(process.get('decision_ready_threshold'))}",
            f"是否可进入正式决策：{'可以' if process.get('decision_ready') else '还不可以'}",
        ]
        for line in process_lines:
            story.append(paragraph(line, body_style))
        round_rows = [["轮次", "阶段", "起点概率", "更新后概率", "准备度", "中间判断"]]
        for item in process.get("rounds", []):
            round_rows.append(
                [
                    str(item["round"]),
                    item["stage"],
                    fmt_pct(item.get("prior_probability_before")),
                    fmt_pct(item.get("posterior_probability_after")),
                    fmt_pct(item.get("decision_readiness")),
                    item.get("interim_judgment") or "-",
                ]
            )
        story.append(make_pdf_table(round_rows, [10 * mm, 24 * mm, 18 * mm, 18 * mm, 18 * mm, 74 * mm]))
        story.append(Spacer(1, 3 * mm))

    story.append(paragraph("决策问题", heading_style))
    question_lines = [
        f"决策问题：{report['question']['decision_question']}",
        f"当前现状：{request_text(request.get('current_state'), 'zh')}",
        f"假设：{report['question']['hypothesis']}",
        f"时间范围：{report['question']['time_horizon']}",
        f"决策截止时间：{report['question']['decision_deadline']}",
        f"成功标准：{report['question']['success_metric']}",
        f"领域：{report['question']['domain']}",
    ]
    for line in question_lines:
        story.append(paragraph(line, body_style))

    story.append(paragraph("先验设置", heading_style))
    prior_lines = [
        f"先验概率：{fmt_pct(report['prior']['probability'])}",
        f"先验分布：{report['prior']['distribution']}",
        f"先验区间：{fmt_interval(report['prior']['credible_interval'])}",
        f"先验来源等级：{report['prior']['source_quality']}",
        f"先验信心：{report['prior']['confidence']}",
        f"等效样本量：{fmt_num(report['prior']['equivalent_sample_size'])}",
    ]
    for line in prior_lines:
        story.append(paragraph(line, body_style))
    for source in report["prior"]["source_summary"]:
        story.append(paragraph(f"来源：{source}", body_style))

    story.append(paragraph("证据摘要", heading_style))
    evidence_rows = [["证据", "等级", "方向", "似然比", "依赖折扣", "摘要"]]
    for item in report["evidence"]:
        evidence_rows.append(
            [
                item["name"],
                item["quality"],
                item["direction"],
                fmt_num(item["likelihood_ratio"]),
                fmt_num(item["dependency_discount"]),
                item.get("summary") or "-",
            ]
        )
    story.append(make_pdf_table(evidence_rows, [40 * mm, 12 * mm, 14 * mm, 16 * mm, 18 * mm, 78 * mm]))
    story.append(Spacer(1, 3 * mm))

    story.append(paragraph("贝叶斯更新", heading_style))
    update_lines = [
        f"更新方法：{report['posterior']['method']}",
        f"后验概率：{fmt_pct(report['posterior']['probability'])}",
        f"后验区间：{fmt_interval(report['posterior']['credible_interval'])}",
        report["natural_frequency"]["explanation"],
    ]
    for line in update_lines:
        story.append(paragraph(line, body_style))

    story.append(paragraph("行动比较", heading_style))
    story.append(paragraph(f"推荐行动：{report['decision']['recommended_action']}", body_style))
    story.append(paragraph(report["decision"]["reason"], body_style))
    action_rows = [["行动", "H 为真时效用", "H 为假时效用", "期望值", "行动阈值"]]
    for action in report["decision"]["expected_value_ranking"]:
        action_rows.append(
            [
                action["name"],
                fmt_num(action["utility_if_h_true"]),
                fmt_num(action["utility_if_h_false"]),
                fmt_num(action["expected_value"]),
                fmt_pct(action["action_threshold"]),
            ]
        )
    story.append(make_pdf_table(action_rows, [54 * mm, 28 * mm, 28 * mm, 24 * mm, 24 * mm]))
    story.append(Spacer(1, 3 * mm))

    story.append(paragraph("敏感性分析", heading_style))
    story.append(paragraph(f"后验范围：{fmt_interval(report['sensitivity']['posterior_range'])}", body_style))
    story.append(paragraph(f"结论稳定性：{report['sensitivity']['conclusion_stability']}", body_style))

    story.append(paragraph("下一步最有价值的信息", heading_style))
    story.append(paragraph(report["next_information"]["recommended_next_information"], body_style))
    story.append(paragraph(report["next_information"]["reason"], body_style))

    story.append(paragraph("Skill 流程", heading_style))
    for index, item in enumerate(SKILL_FLOW, start=1):
        story.append(paragraph(f"{index}. {item['zh']}", body_style))

    story.append(paragraph("Skill 能力", heading_style))
    for item in SKILL_CAPABILITIES:
        story.append(paragraph(f"{item['zh_title']}：{item['zh_body']}", body_style))

    story.append(paragraph("风险与注意事项", heading_style))
    for warning in report["warnings"] or ["当前没有额外警示项。"]:
        story.append(paragraph(f"• {warning}", body_style))

    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        leftMargin=16 * mm,
        rightMargin=16 * mm,
        topMargin=16 * mm,
        bottomMargin=16 * mm,
    )
    doc.build(story)


def configure_docx_font(style, font_name: str, size_pt: float) -> None:
    style.font.name = font_name
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    style.font.size = Pt(size_pt)


def build_docx(request: dict, report: dict, output_path: Path) -> None:
    plain = plain_language_pack(request, report)
    process = report.get("conversation_process") or {}
    document = Document()
    configure_docx_font(document.styles["Normal"], "PingFang SC", 10.5)
    configure_docx_font(document.styles["Title"], "PingFang SC", 20)
    configure_docx_font(document.styles["Heading 1"], "PingFang SC", 16)
    configure_docx_font(document.styles["Heading 2"], "PingFang SC", 13)

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(f"贝叶斯决策报告：{report['title']}")

    note = document.add_paragraph()
    note.add_run("本报告由 yao-bayesian-skill 自动生成，Word 与 HTML、PDF、JSON 共享同一套结构化计算结果。")

    document.add_heading("先说结论", level=1)
    document.add_paragraph(report["summary"]["one_sentence"])
    document.add_paragraph(plain["summary_zh"])

    document.add_heading("你现在该怎么做", level=1)
    for index, step in enumerate(plain["action_steps_zh"], start=1):
        document.add_paragraph(f"{index}. {step}", style="Normal")

    document.add_heading("这份建议成立的前提", level=1)
    for item in plain["premises_zh"]:
        document.add_paragraph(item, style="Normal")
    document.add_paragraph(plain["decision_gate_zh"], style="Normal")

    document.add_heading("为什么不是另外两个选项", level=1)
    if plain["alternatives"]:
        for item in plain["alternatives"]:
            document.add_paragraph(f"{item['name_zh']}：{item['reason_zh']}", style="Normal")
    else:
        document.add_paragraph("当前没有足够的备选行动用于比较。", style="Normal")

    if process:
        document.add_heading("多轮对话过程与决策准备度", level=1)
        document.add_paragraph(process.get("analysis") or "-", style="Normal")
        for line in [
            f"当前状态：{process.get('status')}",
            f"决策准备度：{fmt_pct(process.get('final_readiness'))}",
            f"决策阈值：{fmt_pct(process.get('decision_ready_threshold'))}",
            f"是否可进入正式决策：{'可以' if process.get('decision_ready') else '还不可以'}",
        ]:
            document.add_paragraph(line, style="Normal")
        round_table = document.add_table(rows=1, cols=6)
        round_table.style = "Table Grid"
        for cell, header in zip(round_table.rows[0].cells, ["轮次", "阶段", "起点概率", "更新后概率", "准备度", "中间判断"]):
            cell.text = header
        for item in process.get("rounds", []):
            cells = round_table.add_row().cells
            cells[0].text = str(item["round"])
            cells[1].text = item["stage"]
            cells[2].text = fmt_pct(item.get("prior_probability_before"))
            cells[3].text = fmt_pct(item.get("posterior_probability_after"))
            cells[4].text = fmt_pct(item.get("decision_readiness"))
            cells[5].text = item.get("interim_judgment") or "-"

    document.add_heading("决策问题", level=1)
    for line in [
        f"决策问题：{report['question']['decision_question']}",
        f"当前现状：{request_text(request.get('current_state'), 'zh')}",
        f"假设：{report['question']['hypothesis']}",
        f"时间范围：{report['question']['time_horizon']}",
        f"决策截止时间：{report['question']['decision_deadline']}",
        f"成功标准：{report['question']['success_metric']}",
        f"领域：{report['question']['domain']}",
    ]:
        document.add_paragraph(line, style="Normal")

    document.add_heading("先验设置", level=1)
    for line in [
        f"先验概率：{fmt_pct(report['prior']['probability'])}",
        f"先验分布：{report['prior']['distribution']}",
        f"先验区间：{fmt_interval(report['prior']['credible_interval'])}",
        f"先验来源等级：{report['prior']['source_quality']}",
        f"先验信心：{report['prior']['confidence']}",
        f"等效样本量：{fmt_num(report['prior']['equivalent_sample_size'])}",
    ]:
        document.add_paragraph(line, style="Normal")
    for source in report["prior"]["source_summary"]:
        document.add_paragraph(f"来源：{source}", style="Normal")

    document.add_heading("证据摘要", level=1)
    evidence_table = document.add_table(rows=1, cols=6)
    evidence_table.style = "Table Grid"
    headers = ["证据", "等级", "方向", "似然比", "依赖折扣", "摘要"]
    for cell, header in zip(evidence_table.rows[0].cells, headers):
        cell.text = header
    for item in report["evidence"]:
        cells = evidence_table.add_row().cells
        cells[0].text = item["name"]
        cells[1].text = item["quality"]
        cells[2].text = item["direction"]
        cells[3].text = fmt_num(item["likelihood_ratio"])
        cells[4].text = fmt_num(item["dependency_discount"])
        cells[5].text = item.get("summary") or "-"

    document.add_heading("贝叶斯更新", level=1)
    for line in [
        f"更新方法：{report['posterior']['method']}",
        f"后验概率：{fmt_pct(report['posterior']['probability'])}",
        f"后验区间：{fmt_interval(report['posterior']['credible_interval'])}",
        report["natural_frequency"]["explanation"],
    ]:
        document.add_paragraph(line, style="Normal")

    document.add_heading("行动比较", level=1)
    document.add_paragraph(f"推荐行动：{report['decision']['recommended_action']}", style="Normal")
    document.add_paragraph(report["decision"]["reason"], style="Normal")
    action_table = document.add_table(rows=1, cols=5)
    action_table.style = "Table Grid"
    for cell, header in zip(action_table.rows[0].cells, ["行动", "H 为真时效用", "H 为假时效用", "期望值", "行动阈值"]):
        cell.text = header
    for item in report["decision"]["expected_value_ranking"]:
        cells = action_table.add_row().cells
        cells[0].text = item["name"]
        cells[1].text = fmt_num(item["utility_if_h_true"])
        cells[2].text = fmt_num(item["utility_if_h_false"])
        cells[3].text = fmt_num(item["expected_value"])
        cells[4].text = fmt_pct(item["action_threshold"])

    document.add_heading("敏感性分析", level=1)
    document.add_paragraph(f"后验范围：{fmt_interval(report['sensitivity']['posterior_range'])}", style="Normal")
    document.add_paragraph(f"结论稳定性：{report['sensitivity']['conclusion_stability']}", style="Normal")

    document.add_heading("下一步最有价值的信息", level=1)
    document.add_paragraph(report["next_information"]["recommended_next_information"], style="Normal")
    document.add_paragraph(report["next_information"]["reason"], style="Normal")

    document.add_heading("Skill 流程", level=1)
    for index, item in enumerate(SKILL_FLOW, start=1):
        document.add_paragraph(f"{index}. {item['zh']}", style="Normal")

    document.add_heading("Skill 能力", level=1)
    for item in SKILL_CAPABILITIES:
        document.add_paragraph(f"{item['zh_title']}：{item['zh_body']}", style="Normal")

    document.add_heading("风险与注意事项", level=1)
    for warning in report["warnings"] or ["当前没有额外警示项。"]:
        document.add_paragraph(warning, style="Normal")

    document.save(str(output_path))


def generate_bundle(input_path: Path, output_dir: Path, basename: str | None) -> dict[str, str]:
    request = load_request(input_path)
    report = build_report(request)
    output_dir.mkdir(parents=True, exist_ok=True)
    stem = basename or input_path.stem.replace("_", "-")

    json_path = output_dir / f"{stem}.json"
    md_path = output_dir / f"{stem}.md"
    html_path = output_dir / f"{stem}.html"
    pdf_path = output_dir / f"{stem}.pdf"
    docx_path = output_dir / f"{stem}.docx"

    json_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    md_path.write_text(build_markdown(request, report), encoding="utf-8")
    html_path.write_text(build_html(request, report), encoding="utf-8")
    build_pdf(request, report, pdf_path)
    build_docx(request, report, docx_path)

    return {
        "json": str(json_path),
        "md": str(md_path),
        "html": str(html_path),
        "pdf": str(pdf_path),
        "docx": str(docx_path),
    }


def main() -> None:
    parser = argparse.ArgumentParser(description="Generate a synchronized Bayesian decision report bundle.")
    parser.add_argument("input_json", help="Path to the structured decision request JSON file.")
    parser.add_argument("output_dir", help="Directory where the report bundle should be written.")
    parser.add_argument("--basename", help="Optional basename for output files. Defaults to the input filename stem.")
    args = parser.parse_args()

    generated = generate_bundle(Path(args.input_json), Path(args.output_dir), args.basename)
    print(json.dumps({"ok": True, "generated": generated}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
